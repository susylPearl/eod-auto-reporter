"""
Config Store — persists user configuration as JSON.

Stores settings in ~/Library/Application Support/EOD Reporter/config.json
on macOS. Falls back to ~/.eod-reporter/config.json on other platforms.

Tokens are base64-encoded (not encrypted) for light obfuscation.
"""

from __future__ import annotations

import base64
import json
import os
import platform
from pathlib import Path
from typing import Any, Dict, Optional

_APP_NAME = "EOD Reporter"


def _get_config_dir() -> Path:
    """Return the platform-appropriate config directory."""
    system = platform.system()
    if system == "Darwin":
        base = Path.home() / "Library" / "Application Support"
    elif system == "Windows":
        base = Path(os.environ.get("APPDATA", Path.home() / "AppData" / "Roaming"))
    else:
        base = Path.home() / ".config"
    return base / _APP_NAME


def _config_path() -> Path:
    """Return the full path to config.json."""
    return _get_config_dir() / "config.json"


# Fields that contain sensitive tokens — stored base64-encoded
_SENSITIVE_FIELDS = {"github_token", "clickup_api_token", "slack_bot_token", "openai_api_key"}

# Default configuration values
DEFAULT_CONFIG: Dict[str, Any] = {
    "github_token": "",
    "github_username": "",
    "clickup_api_token": "",
    "clickup_team_id": "",
    "clickup_user_id": "",
    "slack_bot_token": "",
    "slack_channel": "",
    "slack_user_id": "",
    "slack_display_name": "",
    "slack_icon_url": "",
    "slack_monitor_channels": "",
    "report_hour": 18,
    "report_minute": 0,
    "timezone": "Asia/Kathmandu",
    "log_level": "INFO",
    "app_env": "development",
    # AI summarization
    "openai_api_key": "",
    "ai_model": "gemini-2.0-flash",
    "ai_base_url": "https://generativelanguage.googleapis.com/v1beta/openai/",
    "ai_provider": "Google Gemini (Free)",
    # Activity tab customization
    "show_github": True,
    "show_clickup": True,
    "show_slack": True,
    "show_manual": True,
    "show_ai_summary": True,
    "activity_section_order": "github_first",  # or "clickup_first"
    "manual_updates": [],
    "max_commits_per_repo": 10,
    "max_tasks_display": 20,
    "max_comments_display": 10,
}


def _encode(value: str) -> str:
    """Base64-encode a string for light obfuscation."""
    return base64.b64encode(value.encode("utf-8")).decode("ascii")


def _decode(value: str) -> str:
    """Decode a base64-encoded string."""
    try:
        return base64.b64decode(value.encode("ascii")).decode("utf-8")
    except Exception:
        return value  # return as-is if not valid base64


def load_config() -> Dict[str, Any]:
    """
    Load configuration from disk.

    Returns DEFAULT_CONFIG merged with any saved values.
    Missing keys get default values; extra keys are preserved.
    """
    config = dict(DEFAULT_CONFIG)
    path = _config_path()

    if path.exists():
        try:
            raw = json.loads(path.read_text(encoding="utf-8"))
            for key, value in raw.items():
                if key in _SENSITIVE_FIELDS and isinstance(value, str) and value:
                    config[key] = _decode(value)
                else:
                    config[key] = value
        except (json.JSONDecodeError, OSError):
            pass  # corrupt file — use defaults

    return config


def save_config(config: Dict[str, Any]) -> None:
    """
    Persist configuration to disk.

    Sensitive fields are base64-encoded before writing.
    """
    config_dir = _get_config_dir()
    config_dir.mkdir(parents=True, exist_ok=True)

    to_write: Dict[str, Any] = {}
    for key, value in config.items():
        if key in _SENSITIVE_FIELDS and isinstance(value, str) and value:
            to_write[key] = _encode(value)
        else:
            to_write[key] = value

    _config_path().write_text(
        json.dumps(to_write, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )


_ENV_MAP = {
    "GITHUB_TOKEN": "github_token",
    "GITHUB_USERNAME": "github_username",
    "CLICKUP_API_TOKEN": "clickup_api_token",
    "CLICKUP_TEAM_ID": "clickup_team_id",
    "CLICKUP_USER_ID": "clickup_user_id",
    "SLACK_BOT_TOKEN": "slack_bot_token",
    "SLACK_CHANNEL": "slack_channel",
    "SLACK_USER_ID": "slack_user_id",
    "SLACK_DISPLAY_NAME": "slack_display_name",
    "SLACK_ICON_URL": "slack_icon_url",
    "SLACK_MONITOR_CHANNELS": "slack_monitor_channels",
    "REPORT_HOUR": "report_hour",
    "REPORT_MINUTE": "report_minute",
    "TIMEZONE": "timezone",
    "LOG_LEVEL": "log_level",
    "APP_ENV": "app_env",
    "OPENAI_API_KEY": "openai_api_key",
    "AI_MODEL": "ai_model",
    "AI_BASE_URL": "ai_base_url",
    "AI_PROVIDER": "ai_provider",
}

_LABEL_MAP: Dict[str, str] = {
    "github token": "github_token",
    "github username": "github_username",
    "clickup api token": "clickup_api_token",
    "clickup token": "clickup_api_token",
    "clickup team id": "clickup_team_id",
    "team id": "clickup_team_id",
    "clickup user id": "clickup_user_id",
    "slack bot token": "slack_bot_token",
    "bot token": "slack_bot_token",
    "slack channel": "slack_channel",
    "eod channel": "slack_channel",
    "slack user id": "slack_user_id",
    "display name": "slack_display_name",
    "slack display name": "slack_display_name",
    "icon url": "slack_icon_url",
    "slack icon url": "slack_icon_url",
    "monitor channels": "slack_monitor_channels",
    "report hour": "report_hour",
    "hour": "report_hour",
    "report minute": "report_minute",
    "minute": "report_minute",
    "timezone": "timezone",
    "openai api key": "openai_api_key",
    "api key": "openai_api_key",
    "ai model": "ai_model",
    "model": "ai_model",
    "ai base url": "ai_base_url",
    "base url": "ai_base_url",
    "ai provider": "ai_provider",
    "provider": "ai_provider",
}

_INT_FIELDS = {"report_hour", "report_minute"}


def _coerce_value(config_key: str, value: str) -> Any:
    """Cast a raw string value to the right type for a config key."""
    if config_key in _INT_FIELDS:
        try:
            return int(value)
        except ValueError:
            return value
    return value


def _extract_kv_from_text(text: str) -> Dict[str, Any]:
    """
    Extract config values from free-form text by matching known keys.

    Supports:
      - KEY=VALUE  (.env style)
      - "KEY": "VALUE"  (JSON-in-text)
      - Label: value  (human-readable labels)
      - | Label | Value |  (markdown tables)
    """
    import re

    found: Dict[str, Any] = {}

    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue
        # Strip markdown list markers: - , * , + , numbered 1.
        line = re.sub(r'^(?:[-*+]|\d+\.)\s+', '', line).strip()

        # --- KEY=VALUE (.env style) ---
        if "=" in line and not line.startswith("|"):
            key, _, value = line.partition("=")
            key = key.strip().strip('"').strip("'")
            value = value.strip().strip('"').strip("'")
            if key in _ENV_MAP and value:
                cfg_key = _ENV_MAP[key]
                found[cfg_key] = _coerce_value(cfg_key, value)
                continue
            lower = key.lower().replace("_", " ").strip()
            if lower in _LABEL_MAP and value:
                cfg_key = _LABEL_MAP[lower]
                found[cfg_key] = _coerce_value(cfg_key, value)
                continue

        # --- "KEY": "VALUE" (JSON fragments in text) ---
        json_match = re.match(
            r'["\']?([A-Z_a-z][A-Za-z_]*)["\']?\s*:\s*["\'](.+?)["\']', line
        )
        if json_match:
            key, value = json_match.group(1).strip(), json_match.group(2).strip()
            if key in _ENV_MAP and value:
                cfg_key = _ENV_MAP[key]
                found[cfg_key] = _coerce_value(cfg_key, value)
                continue
            lower_key = key.lower().replace("_", " ").strip()
            if lower_key in _LABEL_MAP and value:
                cfg_key = _LABEL_MAP[lower_key]
                found[cfg_key] = _coerce_value(cfg_key, value)
                continue

        # --- Table row: | Label | Value | or Label | Value ---
        if "|" in line:
            cells = [c.strip() for c in line.split("|")]
            cells = [c for c in cells if c and not re.match(r'^[-:]+$', c)]
            if len(cells) >= 2:
                label = cells[0].strip().strip("`*_")
                value = cells[1].strip().strip("`*_")
                lower_label = label.lower().replace("_", " ").strip()
                if lower_label in _LABEL_MAP and value:
                    cfg_key = _LABEL_MAP[lower_label]
                    found[cfg_key] = _coerce_value(cfg_key, value)
                    continue
                upper_label = label.replace(" ", "_").upper()
                if upper_label in _ENV_MAP and value:
                    cfg_key = _ENV_MAP[upper_label]
                    found[cfg_key] = _coerce_value(cfg_key, value)
                    continue

        # --- Label: value (colon-separated, human-readable) ---
        colon_match = re.match(r'^([A-Za-z][A-Za-z_ ]{2,30})\s*:\s*(.+)$', line)
        if colon_match:
            label = colon_match.group(1).strip()
            value = colon_match.group(2).strip().strip("`*_\"'")
            lower_label = label.lower().replace("_", " ").strip()
            if lower_label in _LABEL_MAP and value:
                cfg_key = _LABEL_MAP[lower_label]
                found[cfg_key] = _coerce_value(cfg_key, value)
                continue

    return found


def _read_docx_text(path: Path) -> Optional[str]:
    """Extract plain text from a .docx file. Returns None if unreadable."""
    try:
        from docx import Document
        doc = Document(str(path))
        lines: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append(" | ".join(cells))
        return "\n".join(lines)
    except ImportError:
        return None
    except Exception:
        return None


def import_from_file(file_path: str | Path) -> Optional[Dict[str, Any]]:
    """
    Import settings from a file of any supported format.

    Supported: .env, .txt, .json, .md, .docx, .doc
    Returns a config dict (merged with defaults), or None on failure.
    """
    path = Path(file_path)
    if not path.exists():
        return None

    suffix = path.suffix.lower()

    # --- JSON ---
    if suffix == ".json":
        return _import_json(path)

    # --- Word documents ---
    if suffix in (".docx", ".doc"):
        text = _read_docx_text(path)
        if not text:
            return None
        found = _extract_kv_from_text(text)
        if not found:
            return None
        result = dict(DEFAULT_CONFIG)
        result.update(found)
        return result

    # --- Text-based: .env, .txt, .md, or anything else ---
    try:
        text = path.read_text(encoding="utf-8")
    except (OSError, UnicodeDecodeError):
        return None

    found = _extract_kv_from_text(text)
    if not found:
        return None
    result = dict(DEFAULT_CONFIG)
    result.update(found)
    return result


def _import_json(path: Path) -> Optional[Dict[str, Any]]:
    """
    Import from a JSON file.

    Handles both config-key format (lowercase) and env-var format (UPPER_CASE).
    """
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return None

    if not isinstance(data, dict):
        return None

    result = dict(DEFAULT_CONFIG)
    for key, value in data.items():
        if isinstance(value, (dict, list)):
            continue
        str_val = str(value).strip()
        if not str_val:
            continue
        # Direct config key match
        if key in DEFAULT_CONFIG:
            result[key] = _coerce_value(key, str_val)
        # UPPER_CASE env-var key
        elif key in _ENV_MAP:
            cfg_key = _ENV_MAP[key]
            result[cfg_key] = _coerce_value(cfg_key, str_val)

    return result


def import_from_dotenv(dotenv_path: str | Path) -> Optional[Dict[str, Any]]:
    """
    Import settings from an existing .env file.

    Returns the imported config dict, or None if the file doesn't exist.
    """
    return import_from_file(dotenv_path)


def config_exists() -> bool:
    """Check if a config file already exists."""
    return _config_path().exists()


def get_config_path() -> Path:
    """Return the config file path (for display purposes)."""
    return _config_path()
